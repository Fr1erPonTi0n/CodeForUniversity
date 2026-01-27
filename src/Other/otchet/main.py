import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_word_with_images(folder_path, output_file="photos.docx"):
    image_files = [f for f in os.listdir(folder_path)
                   if f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp'))]

    image_files.sort()
    doc = Document()

    for i, image_file in enumerate(image_files, start=1):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"{i}")
        run.bold = True
        run.font.size = 24
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        full_path = os.path.join(folder_path, image_file)
        doc.add_picture(full_path, width=Inches(6))

        if i < len(image_files):
            doc.add_page_break()

    doc.save(output_file)
    print(f"Документ сохранен как {output_file} с {len(image_files)} изображениями.")


if __name__ == "__main__":
    script_dir = os.path.dirname(os.path.abspath(__file__))
    photos_dir = os.path.join(script_dir, 'photos')
    output_file = os.path.join(script_dir, "photos.docx")
    if not os.path.exists(photos_dir):
        print("Нету папки с фотками")
    else:
        create_word_with_images(photos_dir, output_file)
